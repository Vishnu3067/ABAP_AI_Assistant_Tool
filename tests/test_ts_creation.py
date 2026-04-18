"""
test_ts_creation.py — Tests for the TS Creation Tool feature.

Tests:
  1. extract_fs_text()        — reads a docx and returns text
  2. extract_placeholders()   — finds all {{name}} tokens in the TS template
  3. render_ts_docx()         — renders a filled docx and verifies it is valid
  4. docx_to_preview_html()   — produces HTML fragment with key elements
  5. ask_ai_fill_placeholders() — AI JSON parsing / fallback behaviour (mocked)
  6. /api/ts-creation/generate  — FastAPI endpoint, real multipart upload
  7. /api/ts-creation/regenerate — refine with feedback (session reuse)
  8. /api/ts-creation/download   — download returns docx bytes

Run from the FASTAPI_AI_Tool root:
  python tests/test_ts_creation.py
"""

import io
import json
import re
import sys
import zipfile
from pathlib import Path
from unittest.mock import patch

# Force UTF-8 output so ✅/❌ display correctly on Windows terminals
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# ── Add project root so core/ imports work ─────────────────────────────────
sys.path.insert(0, str(Path(__file__).parent.parent))

# ── Load .env ───────────────────────────────────────────────────────────────
import truststore
truststore.inject_into_ssl()
from dotenv import load_dotenv
load_dotenv(dotenv_path=Path(__file__).parent.parent / ".env", override=True)

# ── ANSI helpers ─────────────────────────────────────────────────────────────
def ok(msg):   print(f"  [PASS] {msg}")
def fail(msg):
    global FAIL
    FAIL += 1
    print(f"  [FAIL] {msg}")
def sep():     print("-" * 65)

PASS = FAIL = 0


def assert_true(condition, label):
    global PASS, FAIL
    if condition:
        ok(label)
        PASS += 1
    else:
        fail(label)
        FAIL += 1


# ─────────────────────────────────────────────────────────────────────────────
# Paths
# ─────────────────────────────────────────────────────────────────────────────
TEMPLATE_PATH = (
    Path(__file__).parent.parent
    / "data"
    / "TS_WRICEF ID_Description v1.0 - AITemplate 1.docx"
)

# Build a minimal synthetic FS docx in memory for upload / text extraction tests
def make_minimal_fs_docx() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading("Functional Specification", level=1)
    doc.add_heading("1. Background", level=2)
    doc.add_paragraph(
        "WRICEF ID: WR-9001.  Process Area: Finance.  "
        "This enhancement creates a custom class /SHL/CL_FIN_NOTIF to send "
        "email notifications after a goods receipt is posted.  "
        "The BAdi implementation resides in enhancement spot ES_FIN_POST."
    )
    doc.add_heading("2. Solution Overview", level=2)
    doc.add_paragraph(
        "The class must implement interface IF_FIN_NOTIF.  "
        "It exposes method SEND_EMAIL(iv_doc_no TYPE ebeln, iv_amount TYPE waers).  "
        "Error handling uses /SHL/CX_ROOT."
    )
    doc.add_heading("3. Stakeholders", level=2)
    tbl = doc.add_table(rows=2, cols=4)
    tbl.rows[0].cells[0].text = "Name"
    tbl.rows[0].cells[1].text = "Role"
    tbl.rows[0].cells[2].text = "Email"
    tbl.rows[0].cells[3].text = "Comment"
    tbl.rows[1].cells[0].text = "John Doe"
    tbl.rows[1].cells[1].text = "Developer"
    tbl.rows[1].cells[2].text = "j.doe@shell.com"
    tbl.rows[1].cells[3].text = "Author"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


FS_DOCX_BYTES = make_minimal_fs_docx()


# ─────────────────────────────────────────────────────────────────────────────
# 1. extract_fs_text
# ─────────────────────────────────────────────────────────────────────────────
print("\n[1] extract_fs_text")
sep()
from core.ts_creation import extract_fs_text

try:
    text = extract_fs_text(FS_DOCX_BYTES)
    assert_true(isinstance(text, str), "Returns a string")
    assert_true(len(text) > 50,        "Non-trivial content extracted")
    assert_true("WR-9001" in text,     "FS paragraph text preserved")
    assert_true("John Doe" in text,    "Table row text preserved")
    assert_true("j.doe@shell.com" in text, "Table cell content preserved")
    # Heading marker present
    assert_true(any(line.startswith("#") for line in text.splitlines()),
                "Heading markers (#) present")
    print(f"  ℹ️  Extracted {len(text)} characters from synthetic FS")
except Exception as e:
    fail(f"extract_fs_text raised: {e}")
    import traceback; traceback.print_exc()


# ─────────────────────────────────────────────────────────────────────────────
# 2. extract_placeholders
# ─────────────────────────────────────────────────────────────────────────────
print("\n[2] extract_placeholders")
sep()
from core.ts_creation import extract_placeholders

try:
    placeholders = extract_placeholders(TEMPLATE_PATH)
    assert_true(isinstance(placeholders, list),    "Returns a list")
    assert_true(len(placeholders) >= 5,            f"At least 5 placeholders found ({len(placeholders)} found)")
    assert_true("RIEFID" in placeholders,          "{{RIEFID}} detected")
    assert_true("Name" in placeholders,            "{{Name}} detected")
    assert_true("date" in placeholders,            "{{date}} detected")
    assert_true("SolutionOverview" in placeholders,"{{SolutionOverview}} detected")
    print(f"  ℹ️  Found {len(placeholders)} unique placeholders: {placeholders}")
except Exception as e:
    fail(f"extract_placeholders raised: {e}")
    import traceback; traceback.print_exc()


# ─────────────────────────────────────────────────────────────────────────────
# 3. render_ts_docx
# ─────────────────────────────────────────────────────────────────────────────
print("\n[3] render_ts_docx")
sep()
from core.ts_creation import render_ts_docx

try:
    # Provide a minimal value set — unknown placeholders get "N/A" from AI
    minimal_values = {p: f"[TEST_{p}]" for p in placeholders}
    minimal_values["RIEFID"]           = "WR-9001"
    minimal_values["Name"]             = "Test Developer"
    minimal_values["date"]             = "18.04.2026"
    minimal_values["SolutionOverview"] = "Test solution overview."

    docx_bytes = render_ts_docx(minimal_values, TEMPLATE_PATH)
    assert_true(isinstance(docx_bytes, bytes),   "Returns bytes")
    assert_true(len(docx_bytes) > 1000,          "Non-trivial docx size")
    # Verify it is a valid zip (docx = zip)
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            names = z.namelist()
        assert_true("word/document.xml" in names, "Contains word/document.xml")
    except zipfile.BadZipFile:
        fail("Rendered bytes are not a valid zip / docx")

    # Verify placeholder values were substituted (no raw {{...}} left for known ones)
    all_xml = ""
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        for n in z.namelist():
            if n.endswith(".xml"):
                all_xml += z.read(n).decode("utf-8", errors="replace")
    leftover = re.findall(r"\{\{([^}]+)\}\}", re.sub(r"<[^>]+>", "", all_xml))
    assert_true(len(leftover) == 0,
                f"No unfilled {{{{...}}}} tokens remain (found: {leftover[:5]})")
    print(f"  ℹ️  Rendered docx size: {len(docx_bytes):,} bytes")
except Exception as e:
    fail(f"render_ts_docx raised: {e}")
    import traceback; traceback.print_exc()


# ─────────────────────────────────────────────────────────────────────────────
# 4. docx_to_preview_html
# ─────────────────────────────────────────────────────────────────────────────
print("\n[4] docx_to_preview_html")
sep()
from core.ts_creation import docx_to_preview_html

try:
    html = docx_to_preview_html(docx_bytes)
    assert_true(isinstance(html, str), "Returns a string")
    assert_true('<div class="ts-doc-preview">' in html, "Outer wrapper div present")
    assert_true("WR-9001" in html,    "RIEFID value present in HTML")
    assert_true("<table" in html,     "At least one table rendered")
    # TS template uses table-based layout (title/headings are in table cells),
    # so standard <h1>-<h6> tags are not expected. Verify rich structure instead.
    assert_true("ts-doc-p" in html or "ts-doc-cell" in html,
                "Body paragraphs or table cells present in preview")
    # No raw placeholders injected as literals
    assert_true("{{" not in html,     "No raw {{...}} in preview HTML")
    print(f"  ℹ️  Preview HTML: {len(html):,} characters")
except Exception as e:
    fail(f"docx_to_preview_html raised: {e}")
    import traceback; traceback.print_exc()


# ─────────────────────────────────────────────────────────────────────────────
# 5. ask_ai_fill_placeholders — mock AI to avoid real LLM call
# ─────────────────────────────────────────────────────────────────────────────
print("\n[5] ask_ai_fill_placeholders (mocked AI)")
sep()
from core.ts_creation import ask_ai_fill_placeholders

# ── 5a: Valid JSON response ──────────────────────────────────────────────────
sample_keys = ["RIEFID", "Name", "date", "SolutionOverview"]
fake_json = json.dumps({k: f"AI_VALUE_{k}" for k in sample_keys})

with patch("core.ts_creation.call_ai", return_value=fake_json):
    result = ask_ai_fill_placeholders(sample_keys, "some fs text")
    assert_true(isinstance(result, dict), "Returns a dict")
    assert_true(result.get("RIEFID") == "AI_VALUE_RIEFID", "Values filled from AI JSON")
    assert_true(set(result.keys()) == set(sample_keys),    "All requested keys present")
    ok("Valid JSON response parsed correctly")

# ── 5b: Markdown-fenced JSON (AI sometimes wraps) ───────────────────────────
fenced = f"```json\n{fake_json}\n```"
with patch("core.ts_creation.call_ai", return_value=fenced):
    result2 = ask_ai_fill_placeholders(sample_keys, "fs text")
    assert_true(result2.get("RIEFID") == "AI_VALUE_RIEFID",
                "Markdown fences stripped correctly")

# ── 5c: Broken JSON → graceful fallback ─────────────────────────────────────
with patch("core.ts_creation.call_ai", return_value="NOT JSON AT ALL"):
    result3 = ask_ai_fill_placeholders(["RIEFID", "Name"], "fs text")
    assert_true(result3.get("RIEFID") == "[RIEFID]",
                "Broken JSON falls back to [PlaceholderName] values")

# ── 5d: Regeneration with feedback and current values ───────────────────────
updated_json = json.dumps({"RIEFID": "WR-UPDATED", "Name": "Jane Doe"})
with patch("core.ts_creation.call_ai", return_value=updated_json):
    result4 = ask_ai_fill_placeholders(
        ["RIEFID", "Name"],
        "fs text",
        feedback="The RIEFID should be WR-UPDATED",
        current_values={"RIEFID": "WR-9001", "Name": "John Doe"},
    )
    assert_true(result4.get("RIEFID") == "WR-UPDATED",
                "Regeneration updates values from feedback")


# ─────────────────────────────────────────────────────────────────────────────
# 6-8: FastAPI endpoint tests via TestClient
# ─────────────────────────────────────────────────────────────────────────────
print("\n[6-8] FastAPI endpoint tests (TestClient — mocked AI)")
sep()

try:
    from fastapi.testclient import TestClient
    from main import app
    client = TestClient(app, raise_server_exceptions=True)
except Exception as e:
    fail(f"Could not import TestClient or app: {e}")
    print(f"  Skipping endpoint tests.")
    client = None

if client:
    # ── Shared mock AI JSON that covers all placeholders ────────────────────
    def _mock_ai_fill(pholders, fs_text, feedback="", current_values=None):
        base = current_values.copy() if current_values else {}
        for p in pholders:
            if p not in base:
                base[p] = f"MockValue_{p}"
        if feedback and "RIEFID" in base:
            base["RIEFID"] = "WR-REGEN"
        return base

    with patch("routes.ts_creation.ask_ai_fill_placeholders", side_effect=_mock_ai_fill):

        # ── [6] POST /api/ts-creation/generate ─────────────────────────────
        print("\n  [6] POST /api/ts-creation/generate")
        resp = client.post(
            "/api/ts-creation/generate",
            files={"fs_file": ("TestFS.docx", FS_DOCX_BYTES,
                               "application/vnd.openxmlformats-officedocument"
                               ".wordprocessingml.document")},
        )
        assert_true(resp.status_code == 200,
                    f"HTTP 200 (got {resp.status_code}): {resp.text[:120]}")
        if resp.status_code == 200:
            data = resp.json()
            assert_true("doc_id" in data,           "Response contains doc_id")
            assert_true("preview_html" in data,     "Response contains preview_html")
            assert_true("placeholder_count" in data,"Response contains placeholder_count")
            assert_true(data["placeholder_count"] >= 5,
                        f"Placeholder count ≥ 5 (got {data['placeholder_count']})")
            assert_true('<div class="ts-doc-preview">' in data["preview_html"],
                        "Preview HTML has ts-doc-preview wrapper")
            doc_id = data["doc_id"]
            print(f"    ℹ️ doc_id = {doc_id}")
            print(f"    ℹ️ {data['placeholder_count']} placeholders filled")

        # ── Reject non-docx ──────────────────────────────────────────────────
        bad_resp = client.post(
            "/api/ts-creation/generate",
            files={"fs_file": ("notes.txt", b"hello", "text/plain")},
        )
        assert_true(bad_resp.status_code == 400, "Rejects non-.docx file with 400")

        # ── Reject empty file ────────────────────────────────────────────────
        empty_resp = client.post(
            "/api/ts-creation/generate",
            files={"fs_file": ("empty.docx", b"PK\x03\x04" + b"\x00" * 50,
                               "application/vnd.openxmlformats-officedocument"
                               ".wordprocessingml.document")},
        )
        assert_true(empty_resp.status_code in (400, 500),
                    "Rejects unreadable/empty docx with 4xx/5xx")

    # Regenerate + download tests need a valid session → redo with mock
    with patch("routes.ts_creation.ask_ai_fill_placeholders", side_effect=_mock_ai_fill):
        gen_resp = client.post(
            "/api/ts-creation/generate",
            files={"fs_file": ("TestFS.docx", FS_DOCX_BYTES,
                               "application/vnd.openxmlformats-officedocument"
                               ".wordprocessingml.document")},
        )
        doc_id = gen_resp.json()["doc_id"] if gen_resp.status_code == 200 else None

        # ── [7] POST /api/ts-creation/regenerate ───────────────────────────
        print("\n  [7] POST /api/ts-creation/regenerate")
        if doc_id:
            regen_resp = client.post(
                "/api/ts-creation/regenerate",
                json={"doc_id": doc_id, "feedback": "The RIEFID is wrong, set it to WR-REGEN"},
            )
            assert_true(regen_resp.status_code == 200,
                        f"HTTP 200 (got {regen_resp.status_code}): {regen_resp.text[:120]}")
            if regen_resp.status_code == 200:
                rdata = regen_resp.json()
                assert_true("doc_id" in rdata,       "Regen response has doc_id")
                assert_true("preview_html" in rdata,  "Regen response has preview_html")
                assert_true(rdata["doc_id"] == doc_id,"doc_id unchanged after regen")

            # Blank feedback → 400
            blank_resp = client.post(
                "/api/ts-creation/regenerate",
                json={"doc_id": doc_id, "feedback": "   "},
            )
            assert_true(blank_resp.status_code == 400, "Blank feedback returns 400")

            # Unknown doc_id → 404
            bad_regen = client.post(
                "/api/ts-creation/regenerate",
                json={"doc_id": "00000000-dead-beef-0000-000000000000", "feedback": "fix it"},
            )
            assert_true(bad_regen.status_code == 404, "Unknown doc_id returns 404")
        else:
            fail("Skipped regen tests — generate did not return doc_id")

        # ── [8] GET /api/ts-creation/download/{doc_id} ─────────────────────
        print("\n  [8] GET /api/ts-creation/download/{doc_id}")
        if doc_id:
            dl_resp = client.get(f"/api/ts-creation/download/{doc_id}")
            assert_true(dl_resp.status_code == 200,
                        f"HTTP 200 (got {dl_resp.status_code}): {dl_resp.text[:120]}")
            if dl_resp.status_code == 200:
                ct = dl_resp.headers.get("content-type", "")
                cd = dl_resp.headers.get("content-disposition", "")
                assert_true("wordprocessingml" in ct, f"Content-Type is docx ({ct})")
                assert_true("attachment" in cd,        f"Content-Disposition is attachment ({cd})")
                assert_true(".docx" in cd,             f"Filename ends with .docx ({cd})")
                # Verify the bytes are a valid zip / docx
                try:
                    zipfile.ZipFile(io.BytesIO(dl_resp.content))
                    ok("Downloaded file is a valid .docx (zip)")
                    PASS += 1
                except zipfile.BadZipFile:
                    fail("Downloaded file is NOT a valid zip")
                    FAIL += 1
                print(f"    ℹ️ Downloaded {len(dl_resp.content):,} bytes")

            # Unknown id → 404
            bad_dl = client.get("/api/ts-creation/download/00000000-dead-beef-0000-000000000000")
            assert_true(bad_dl.status_code == 404, "Unknown doc_id download returns 404")
        else:
            fail("Skipped download tests — no doc_id")


# ─────────────────────────────────────────────────────────────────────────────
# Summary
# ─────────────────────────────────────────────────────────────────────────────
print()
sep()
total = PASS + FAIL
print(f"  Results: {PASS}/{total} passed,  {FAIL} failed")
sep()
sys.exit(0 if FAIL == 0 else 1)
